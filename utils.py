import requests
import json
import urllib
import os

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor, Pt, Inches
from docx.image.exceptions import UnrecognizedImageError

from PIL import Image

from gitutils import add_hyperlink

# General API Requests Function


def getJsonData(api_url, team_slug):
    complete_url = api_url + team_slug
    response = requests.get(complete_url)
    data = response.json()
    json = data

    return json


def writeTeamNameHeader(team_info_json, paragraph):
    institution = team_info_json.get('name')
    pg = paragraph.clear()
    pg.add_run("Overview ").bold = True
    pg.add_run("and ")
    pg.add_run("Executive Summary ").bold = True
    pg.add_run("of ")
    pg.add_run(institution)

# Currently Ignores GaTech for 2016


def getTeamSlugs(year):
    try:
        list_of_slugs = []
        api_url = "https://nai.nasa.gov/api/teams/teams/?year=" + str(year)
        response = requests.get(api_url)
        team_list = response.json()
        for team in team_list:
            if team.get('slug') == "gatech" and year == 2016:
                pass
            else:
                list_of_slugs.append(team.get('slug'))
        return list_of_slugs
    except:
        print response
        exit(1)

### Team Members Function ###

# Adding Pillow image reduction - JH 12/1/16
def getAvatar(team_slug, person):
    try:
        first = person.get('first_name').lower()
        last = person.get('last_name').lower()
        int_name = team_slug
        cached_name = int_name + "_" + first + "_" + last + ".jpg"
        if person.get('avatars')[0]:
            image_url = "http://nai.nasa.gov" + person.get('avatars')[0][0]
            urllib.urlretrieve(image_url, cached_name)
        target = Image.open(cached_name)
        target = target.resize((160, 300), Image.ANTIALIAS)
        new_cache_name = int_name + "_" + first + "_" + last + "_" + "scaled" + ".jpg"
        target.save(new_cache_name, optimize=True, quality=90)
        return new_cache_name
    except IndexError:
        # No Image Found. Replaced with Generic Astronaut Image.
        cached_name = "no_avatar.png"
        return cached_name
    except IOError:
        first = person.get('first_name').lower()
        last = person.get('last_name').lower()
        int_name = team_slug
        cached_name = int_name + "_" + first + "_" + last + ".jpg"
        if person.get('avatars')[0]:
            image_url = "http://nai.nasa.gov" + person.get('avatars')[0][0]
            urllib.urlretrieve(image_url, cached_name)
        target = Image.open(cached_name).convert('RGB')
        target = target.resize((160, 300), Image.ANTIALIAS)
        new_cache_name = int_name + "_" + first + "_" + last + "_" + "scaled" + ".jpg"
        target.save(new_cache_name, optimize=True, quality=90)
        return new_cache_name



def writePeopleInfoToWord(people, team_slug, document):
    p = document.add_paragraph()
    p.add_run(
        'Member Roster\n\t- Please identify students as Undergrad, Grad, or PhD').bold = True
    p.add_run(
        '\n\t- If adding an image provide it at 600 x 600px or 2" x 2" 300dpi').bold = True
    p.add_run().add_break()
    for index, person in enumerate(people):
        # Pull all the info out and store avatar image locally.
        avatar_image_name = getAvatar(team_slug, person)
        first = person.get('first_name')
        last = person.get('last_name')
        institution = person.get('institution')
        email = person.get('email')
        if email is None:
            email = " "

        try:
            table = document.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            pg = hdr_cells[0].add_paragraph()
            r = pg.add_run()
            r.add_picture(avatar_image_name, width=Inches(1.00), height=Inches(1.00))
            hdr_cells[1].text = first
            hdr_cells[2].text = last
            hdr_cells[3].text = institution
            hdr_cells[4].merge(hdr_cells[5])
            hdr_cells[4].text = email

            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].merge(row_cells[2])
            row_cells[3].merge(row_cells[4])
            row_cells[3].merge(row_cells[5])

            # Assumption: PI is index 0.
            if index == 0:
                row_cells[0].text = "[  ] Was this person's info updated?"
                row_cells[
                    3].text = " "
            else:
                row_cells[0].text = "[  ] Was this person's info updated?\n[  ] Under\t[  ] Grad\t[  ] PhD"
                row_cells[
                    3].text = "[  ] Would you like us to remove this member from your team?"
                row_cells[3].bold = True

        except (ZeroDivisionError):
            print team_slug
            print first
            print last
            print institution
            print email
            print avatar_image_name
            print "Image Pixel Error\n"
            pass

        except TypeError:
            print team_slug
            print first
            print last
            print institution
            print email
            print avatar_image_name
            print "None Type Error\n"
            pass

        # Unknown Error that occurs with image stream via python-docx. - JH 1/12/17
        except UnrecognizedImageError:
            print team_slug
            print first
            print last
            print institution
            print email
            print avatar_image_name
            print "Unrecognized Image Error\n"
            pass


### Publications Function ###
def writePubsToWord(publications, document):
    p = document.add_paragraph()
    p.add_run('Publications').bold = True
    p.add_run().add_break()
    for month_data in publications:
        for pub_data in month_data.get('publication_set'):
            if pub_data.get('doi'):
                doi_string = "DOI: " + pub_data.get('doi')
                raw_submission = pub_data.get(
                    'raw_submission').replace('\r\n', ' ')
                write_to_doc_string = "[  ] " + \
                    doi_string + "\n" + raw_submission + "\n"
            else:
                doi_string = "DOI: " + "None"
                raw_submission = pub_data.get(
                    'raw_submission').replace('\r\n', ' ')
                write_to_doc_string = "[  ] " + \
                    doi_string + "\n" + raw_submission + "\n"
            document.add_paragraph(write_to_doc_string, style="ListBullet")

    return document

### Logo Function ###


def getLogo(logo_url, team_slug, size):
    if logo_url:
        cached_name = team_slug + "_" + size.lower() + ".jpg"
        urllib.urlretrieve(logo_url, cached_name)
    return cached_name


def writeLogosToWord(logos, team_slug, document):
    p = document.add_paragraph()
    p.add_run("Current Logos").bold = True
    p.add_run().add_break()

    table = document.add_table(rows=1, cols=1, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "[  ] Would you like us to change your team logo?"

    logo_image_names = []
    for image in logos:
        logo_image_name = getLogo(
            image.get('logo'), team_slug, image.get('size'))
        logo_image_names.append(logo_image_name)

    row_cells = table.add_row().cells
    pg = row_cells[0].add_paragraph()
    r = pg.add_run()
    r.add_picture(logo_image_names[0])

    return document

# Write Interdisciplinarity to Section


def writeIDsToWord(overview, paragraph):
    nai_url = "https://nai.nasa.gov"
    project_report_text = overview.get('naiprojectreport_set')
    pg = paragraph.clear()
    pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pg.add_run("\n")
    for item in project_report_text:
        if item.get('interdisciplinarity'):
            if item.get("title") and item.get("get_absolute_url"):
                proj_url = nai_url + item.get("get_absolute_url")
                pg.add_run("\n")
                pg.add_run("Title: ").bold = True
                add_hyperlink(pg, proj_url, item.get("title"), '0000FF', True)
                pg.add_run("\n")
                pg.add_run("\n- ").bold = True
                pg.add_run(item.get('interdisciplinarity') + "\n")
            else:
                pg.add_run("\n")
                pg.add_run("Title: ").bold = True
                pg.add_run(item.get("title"))
                pg.add_run("\n")
                pg.add_run("\n- ").bold = True
                pg.add_run(item.get('interdisciplinarity') + "\n")

# Write Collaborations to Section


def writeCollabsToWord(overview, paragraph):
    nai_url = "https://nai.nasa.gov"
    project_report_text = overview.get('naiprojectreport_set')
    pg = paragraph.clear()
    pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pg.add_run("\n")
    for item in project_report_text:
        if item.get('cross_team_collaboration'):
            if item.get("title") and item.get("get_absolute_url"):
                proj_url = nai_url + item.get("get_absolute_url")
                pg.add_run("\n")
                pg.add_run("Title: ").bold = True
                add_hyperlink(pg, proj_url, item.get("title"), '0000FF', True)
                pg.add_run("\n")
                pg.add_run("\n- ").bold = True
                pg.add_run(item.get('cross_team_collaboration') + "\n\n")
            else:
                pg.add_run("\n")
                pg.add_run("Title: ").bold = True
                pg.add_run(item.get("title"))
                pg.add_run("\n")
                pg.add_run("\n- ").bold = True
                pg.add_run(item.get('interdisciplinarity') + "\n")

# Write Field Site Data to Section


def getFieldSiteData(fs):
    if fs.get("name"):
        name = fs.get("name")
    if fs.get("point"):
        point = fs.get("point")
        point_list = point.split(" ")
        latitude = point_list[2].rstrip(")")
        longitude = point_list[1].lstrip("(")

    if fs.get("description"):
        if fs.get("description") == "&nbsp;":
            description = "N/A"
        else:
            description = fs.get("description")
    else:
        description = "N/A"
    if fs.get("characterize_research"):
        if fs.get("characterize_research") == "&nbsp;":
            cr = "N/A"
        else:
            cr = fs.get("characterize_research")
    else:
        cr = "N/A"

    fs_data = [name, latitude, longitude]

    return fs_data, description, cr


def writeFieldSiteDataToWord(overview, doc):
    project_report_text = overview.get('naiprojectreport_set')
    doc.add_page_break()
    pg = doc.add_paragraph()
    for item in project_report_text:
        if item.get("field_sites"):
            for fs in item.get("field_sites"):
                if "Observatory" in fs.get("name") or "Telescope" in fs.get("name"):
                    pass
                else:
                    fs_list, description, cr = getFieldSiteData(fs)
                    pg.add_run("Location Name:\n")
                    pg.add_run(fs_list[0] + "\n\n")
                    pg.add_run("Latitude and Longitude: " +
                           fs_list[1] + ", " + fs_list[2])
                    pg.add_run("\n\n")
                    pg.add_run("One sentence description:")
                    pg.add_run("\n\n")

                    if description != "N/A":
                        pg.add_run("Description: ").bold = True
                        pg.add_run(description)
                        pg.add_run("\n\n")
                    if cr != "N/A":
                        pg.add_run("Characterize Research: ").bold = True
                        pg.add_run(cr)
                        pg.add_run("\n\n")

                    pg.add_run("Category tag: ")

                    red_text = pg.add_run("select one\n")
                    red_text.font.color.rgb = RGBColor(255, 0, 0)

                    pg.add_run("\t[ ] Ancient Earth sites\n")
                    pg.add_run("\t[ ] Banded Iron Formations\n")
                    pg.add_run("\t[ ] Deep Subsurface sites\n")
                    pg.add_run("\t[ ] Deserts\n")
                    pg.add_run("\t[ ] Extreme Lakes\n")
                    pg.add_run("\t[ ] Extremophiles\n")
                    pg.add_run("\t[ ] Highly Acidic Sites\n")
                    pg.add_run("\t[ ] Hydrothermal Systems\n")
                    pg.add_run("\t[ ] Serpentinization sites\n")
                    pg.add_run("\n\n")
                    pg.add_run("Overview: ")
                    pg.add_run("\n\n")
                    pg.add_run("Astrobiology research conducted at the site:")
                    pg.add_run("\n\n")
                    pg.add_run("Related links:")
                    run = pg.add_run("\n\n")
                    run.add_break(WD_BREAK.PAGE)
    doc.add_page_break()

# Write Team Overview Set to Word Doc


def writeTeamOverviewToWord(overview, paragraph):
    overview_text = overview.get('overview')
    split_text = overview_text.split("\r\n")
    if "Pre-populate" in paragraph.text:
        pg = paragraph.clear()
    pg = paragraph
    for index, text in enumerate(split_text):
        if index is 0:
            pg.add_run(text)
        elif index == (len(split_text) - 1):
            pg.add_run("\n\t" + text + "\n\n")
        else:
            pg.add_run("\n\t" + text + "\n")

# Seperate Bold Text from Associated Body Paragraph


def seperateTitleAndBody(body_string):
    split_text_list = body_string.split("</b>\r\n")
    cleaned_up_title = split_text_list[0].lstrip("<b>")
    cleaned_up_body = split_text_list[1].replace("\r\n", " ")

    return cleaned_up_title, cleaned_up_body


def writeExecSummaryToWord(exec_summary, paragraph):
    exec_sum = exec_summary.get("naiexecutivesummary_set")[0]
    exec_text = exec_sum.get('text')
    exec_pg_list = exec_text.split("\r\n\r\n")
    pg = paragraph.clear()
    pg.add_run("\n")
    for text in exec_pg_list:
        if "<b>" in text:
            title, body = seperateTitleAndBody(text)
            pg.add_run(title).bold = True
            pg.add_run("\n")
            pg.add_run(body + "\n\n")
        else:
            cleaned_up_text = text.replace("\r\n", " ")
            pg.add_run(cleaned_up_text + "\n\n")
