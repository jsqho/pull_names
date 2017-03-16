import requests
import json
import os
import urllib
import sys

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor, Pt

from utils import (getJsonData, getTeamSlugs,
                   writePeopleInfoToWord, writePubsToWord,
                   writeLogosToWord, writeTeamOverviewToWord,
                   writeExecSummaryToWord, writeTeamNameHeader,
                   writeIDsToWord, writeCollabsToWord, writeFieldSiteDataToWord)


nai_team_api = "https://nai.nasa.gov/api/teams/teams/?year=2015&team="
mr_team_api = "https://nai-monthly.arc.nasa.gov/api/reports/?year=2016&team="
# Prepopulate from previous year's reports
team_project_report_api = "https://nai.nasa.gov/api/reports/team-reports/?year=2015&team="
# Prepopulate from current year's overview text
team_overview_api = "https://nai.nasa.gov/api/reports/team-reports/?year=2016&team="
team_slugs = getTeamSlugs(2016)
# Only for Testing
#team_slugs = ["arc", ]

os.system("mkdir media")
os.system("mkdir reports")
try:
    document_title = "Overview and Executive Summary of NAI Team"
    exec_sum_tag = "this section should be written as a coherent piece independent of the individual project reports"
    interdisciplinarity_tag = "* Interdisciplinary research"
    collab_tag = "Has this project involved collaboration that has benefitted from the NAI?"
    field_site_tag = "Related links: websites, publications, articles, videos"

    for team_slug in team_slugs:
        team_info_json = getJsonData(nai_team_api, team_slug)
        # Due to CAN 5 extension, we need to catch for CAN 6 or 7. - JH 1/12/17
        if len(team_info_json) == 2:
            for team in team_info_json:
                if team.get('can') == 7 or team.get('can') == 6:
                    team_info_json = team
        else:
            team_info_json = team_info_json[0]
        team_project_report_json = getJsonData(
            team_project_report_api, team_slug)[0]
        team_overview_json = getJsonData(team_overview_api, team_slug)[0]
        team_exec_json = getJsonData(team_overview_api, team_slug)[0]
        #team_pubs_data = getJsonData(mr_team_api, team_slug)

        # See L74 - JH 11/28/16
        # logos = team_info_json.get('logos')

        people = team_info_json.get('members')

        doc = Document("blank.docx")
        try:
            styles = doc.styles
            style = styles.add_style('ListBullet', WD_STYLE_TYPE.PARAGRAPH)
            style = styles.add_style('Table Grid', WD_STYLE_TYPE.TABLE)

            title_style = doc.styles
            obj_title = title_style.add_style(
                "DefaultStyle", WD_STYLE_TYPE.CHARACTER)
            title_font = obj_title.font
            title_font.size = Pt(12)

            red_style = doc.styles
            obj_red = red_style.add_style(
                "RedStyle", WD_STYLE_TYPE.CHARACTER)
            red_font = obj_red.font
            red_font = RGBColor(from_string(ff0000))
        except ValueError:
            pass

        ss = len(doc.paragraphs) - 1
        for index, paragraph in enumerate(doc.paragraphs):
            if document_title in paragraph.text:
                #writeTeamNameHeader(team_info_json, paragraph)
                pass
            elif "Pre-populate" in paragraph.text:
                # Cannot return document or else paragraph is written at bottom of
                # document instead. - JH
                # Removed Project Reports Via Julie's Request. - JH 11/29/16
                # Interdisciplinarity and Collaborations will be moved to appropriate sections.
                # writeProjectReportsToWord(team_project_report_json,
                # paragraph) ###
                #writeTeamOverviewToWord(team_overview_json, paragraph)
                pass
            elif exec_sum_tag in paragraph.text:
                # Locate Placeholder Text, go 2 indices ahead, and write Executive
                # Summaries there.
                # writeExecSummaryToWord(
                #    team_exec_json, doc.paragraphs[index + 2])
                pass
            elif interdisciplinarity_tag in paragraph.text:
                # writeIDsToWord(team_project_report_json,
                #               doc.paragraphs[index + 2])
                pass
            elif collab_tag in paragraph.text:
                # writeCollabsToWord(team_project_report_json,
                #                   doc.paragraphs[index + 2])
                pass
            elif field_site_tag in paragraph.text:
                #writeFieldSiteDataToWord(team_project_report_json, doc)
                pass
            elif index == ss:
                # Removing Logos and Pubs. They will add one themselves in this report. - JH 11/29/16
                # Pubs will be in seperate document. - JH 11/29/16
                # new_doc = writePubsToWord(team_pubs_data, doc) ###
                # new_doc.add_page_break() ###
                # new_doc = writeLogosToWord(logos, team_slug, doc) ###
                # new_doc.add_page_break() ###
                writePeopleInfoToWord(people, team_slug, doc)

        doc_name = team_slug + "_updated_people"
        doc.save("%s.docx" % doc_name)
        os.system("mv %s.docx reports/" % doc_name)
        os.system("mv *.jpg media/")

except:
    print team_slug
    print sys.exc_info()[0]
    raise
